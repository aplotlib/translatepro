import streamlit as st
import PyPDF2
import docx
import tempfile
import os
from pathlib import Path
import time
from transformers import MarianMTModel, MarianTokenizer
import torch
import base64
from io import BytesIO

st.set_page_config(
    page_title="Document Translator",
    page_icon="🌐",
    layout="wide"
)

# Cache the models to avoid reloading them
@st.cache_resource
def load_es_en_model():
    model_name = "Helsinki-NLP/opus-mt-es-en"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    return tokenizer, model

@st.cache_resource
def load_zh_en_model():
    model_name = "Helsinki-NLP/opus-mt-zh-en"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    return tokenizer, model

# Translation function
def translate_text(text, source_lang):
    # Split text into manageable chunks (about 100 words per chunk)
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for line in text.split('\n'):
        words = line.split()
        if current_word_count + len(words) > 100:
            chunks.append(' '.join(current_chunk))
            current_chunk = words
            current_word_count = len(words)
        else:
            current_chunk.extend(words)
            current_word_count += len(words)
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    # Load appropriate model based on source language
    if source_lang == 'Spanish':
        tokenizer, model = load_es_en_model()
    else:  # Chinese
        tokenizer, model = load_zh_en_model()
    
    # Translate each chunk
    translated_chunks = []
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, chunk in enumerate(chunks):
        status_text.text(f"Translating chunk {i+1}/{len(chunks)}")
        
        # Skip empty chunks
        if not chunk.strip():
            translated_chunks.append("")
            progress_bar.progress((i + 1) / len(chunks))
            continue
        
        # Tokenize and translate
        batch = tokenizer([chunk], return_tensors="pt", padding=True)
        
        # Handle device placement
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model = model.to(device)
        batch = {k: v.to(device) for k, v in batch.items()}
        
        # Generate translation with a reasonable max length
        gen_kwargs = {
            "max_length": min(512, len(chunk.split()) * 2),  # Roughly double the input length
            "num_beams": 4,
            "early_stopping": True
        }
        
        # Translate
        with torch.no_grad():
            generated_ids = model.generate(**batch, **gen_kwargs)
        
        # Decode the generated output
        translated_text = tokenizer.batch_decode(generated_ids, skip_special_tokens=True)[0]
        translated_chunks.append(translated_text)
        
        # Update progress
        progress_bar.progress((i + 1) / len(chunks))
    
    status_text.text("Translation complete!")
    return '\n'.join(translated_chunks)

# Functions to handle different file formats
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    total_pages = len(pdf_reader.pages)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, page in enumerate(pdf_reader.pages):
        status_text.text(f"Extracting text from page {i+1}/{total_pages}")
        text += page.extract_text() + "\n\n"
        progress_bar.progress((i + 1) / total_pages)
    
    status_text.text("Text extraction complete!")
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    total_paragraphs = len(doc.paragraphs)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, para in enumerate(doc.paragraphs):
        status_text.text(f"Extracting text from paragraph {i+1}/{total_paragraphs}")
        text += para.text + "\n"
        progress_bar.progress((i + 1) / total_paragraphs)
    
    status_text.text("Text extraction complete!")
    return text

def extract_text_from_txt(file):
    return file.read().decode('utf-8')

def save_docx(translated_text):
    doc = docx.Document()
    for para in translated_text.split('\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Save to BytesIO object
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def get_download_link(content, filename, text):
    b64 = base64.b64encode(content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

# Main UI
st.title("🌐 Document Translator")
st.subheader("Translate documents from Spanish and Mandarin to English")

# Sidebar
st.sidebar.title("Settings")
source_language = st.sidebar.radio("Source Language", ["Spanish", "Chinese (Mandarin)"])
output_format = st.sidebar.radio("Output Format", ["Text", "Word Document (.docx)"])

# File uploader
uploaded_file = st.file_uploader("Upload your document (PDF, DOCX, or TXT)", 
                                type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    st.info(f"File uploaded: {uploaded_file.name}")
    file_extension = Path(uploaded_file.name).suffix.lower()
    
    with st.spinner("Processing..."):
        # Extract text based on file type
        if file_extension == ".pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif file_extension == ".docx":
            text = extract_text_from_docx(uploaded_file)
        elif file_extension == ".txt":
            text = extract_text_from_txt(uploaded_file)
        else:
            st.error("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
            st.stop()
        
        # Display extracted text
        st.subheader("Extracted Text")
        with st.expander("Show original text"):
            st.text_area("Original", text, height=200)
        
        # Translate button
        if st.button("Translate Document"):
            with st.spinner("Translating... This may take several minutes for large documents"):
                start_time = time.time()
                translated_text = translate_text(text, source_language)
                end_time = time.time()
                
                st.success(f"Translation completed in {end_time - start_time:.2f} seconds!")
                
                # Display translated text
                st.subheader("Translated Text")
                st.text_area("Translation", translated_text, height=300)
                
                # Download options
                st.subheader("Download Options")
                
                if output_format == "Text":
                    # Text download
                    text_bytes = translated_text.encode()
                    st.download_button(
                        label="Download as Text",
                        data=text_bytes,
                        file_name="translated_document.txt",
                        mime="text/plain"
                    )
                else:
                    # DOCX download
                    docx_bytes = save_docx(translated_text)
                    st.download_button(
                        label="Download as Word Document",
                        data=docx_bytes,
                        file_name="translated_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

# Instructions and information
with st.expander("How to use this app"):
    st.markdown("""
    1. Select the source language (Spanish or Chinese) from the sidebar
    2. Choose your preferred output format
    3. Upload your document (PDF, DOCX, or TXT format)
    4. Click "Translate Document"
    5. Review the translated text
    6. Download the translation in your preferred format
    
    **Note:** Translation of large documents can take several minutes. The app processes text in chunks to handle documents of any size.
    """)

with st.expander("About this translator"):
    st.markdown("""
    This document translator uses the Marian Neural Machine Translation (MarianNMT) models to translate from 
    Spanish and Chinese to English. It's designed to handle large documents (30+ pages) efficiently.
    
    **Features:**
    - Free and open-source translation
    - Support for PDF, DOCX, and TXT files
    - Word-for-word translation
    - No usage limits
    - Local processing (no data sent to external APIs)
    - Download options for the translated text
    
    The translation quality depends on the complexity of the source text and document formatting.
    """)

# Footer
st.markdown("---")
st.markdown("Powered by Hugging Face's MarianMT models")
